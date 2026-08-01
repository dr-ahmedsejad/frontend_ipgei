"""
Génère un questionnaire d'audit SI au format Word (.docx) — VERSION FLASH.

Une seule question pivot par axe — 14 questions au total.
Diagnostic rapide pour la direction (15-20 min de remplissage).

Sortie : docs/audit_questionnaire_si_flash.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# 14 questions pivots — révélatrices, formulées comme un test concret
# ---------------------------------------------------------------------------

QUESTIONS = [
    (
        1, "Pilotage informatique",
        "Si votre système informatique tombait en panne demain matin, sauriez-vous "
        "nommer en moins de 30 secondes la personne à appeler en première intention ?"
    ),
    (
        2, "Sécurité des données",
        "Si vos serveurs étaient détruits cette nuit (incendie, vol, ransomware), "
        "pourriez-vous reprendre l'activité avec toutes les données intactes en moins de 24 h ?"
    ),
    (
        3, "Traçabilité des actions",
        "Pouvez-vous savoir QUI a modifié la note d'un étudiant, QUAND et POURQUOI, "
        "sans demander à personne ?"
    ),
    (
        4, "Référentiels académiques",
        "Vos filières, modules, salles, semestres et calendriers universitaires "
        "sont-ils gérés dans un seul et même système informatique ?"
    ),
    (
        5, "Pré-inscriptions et admissions",
        "Un candidat peut-il déposer sa candidature 100 % en ligne (formulaire, pièces "
        "jointes, suivi du statut) sans se déplacer à l'établissement ?"
    ),
    (
        6, "Vie étudiante et scolarité",
        "Vos étudiants ont-ils un espace personnel en ligne où ils consultent en "
        "autonomie leurs notes, emploi du temps, absences et documents officiels ?"
    ),
    (
        7, "Emplois du temps",
        "Vos emplois du temps sont-ils générés par un outil dédié qui détecte "
        "automatiquement les conflits (enseignant, salle, groupe), ou encore élaborés sur Excel ?"
    ),
    (
        8, "Enseignants et paiement des vacations",
        "Les états de paiement des vacations sont-ils générés AUTOMATIQUEMENT à partir "
        "des heures effectivement réalisées (pointage), ou recalculés manuellement chaque mois ?"
    ),
    (
        9, "Présences et absences",
        "Pouvez-vous savoir EN UN CLIC combien d'heures un étudiant donné a manqué "
        "ce semestre, avec les justificatifs associés ?"
    ),
    (
        10, "Notes, évaluations et délibérations",
        "Vos procès-verbaux de délibération (moyennes, décisions, rachats, codes statut) "
        "sont-ils générés AUTOMATIQUEMENT par le système, ou rédigés manuellement après chaque session ?"
    ),
    (
        11, "Stages et insertion professionnelle",
        "Le cycle complet d'un stage (convention → suivi → évaluation → soutenance → PFE) "
        "est-il totalement numérisé et tracé dans le système ?"
    ),
    (
        12, "Réclamations étudiantes",
        "Un étudiant peut-il déposer une réclamation (note, absence, autre) EN LIGNE "
        "et suivre le statut de son traitement, sans déplacement physique ?"
    ),
    (
        13, "Documents et reporting",
        "Pouvez-vous éditer une attestation d'inscription, un relevé de notes ou un "
        "tableau de bord d'activité en MOINS D'UNE MINUTE ?"
    ),
    (
        14, "Site web et présence numérique",
        "Si un candidat tape le nom de votre établissement sur Google aujourd'hui, "
        "ce qu'il trouve (site, réseaux, témoignages) lui donne-t-il VRAIMENT envie de s'inscrire ?"
    ),
]


# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

PRIMARY_COLOR = "1F4E79"
SECONDARY_COLOR = "2E75B6"
LIGHT_FILL = "F2F6FA"


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_run_font(run, name="Calibri", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_compact(p, before=0, after=0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

def build_document(out_path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ----- Titre -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(p, after=2)
    r = p.add_run("Diagnostic flash du Système d'Information")
    set_run_font(r, size=15, bold=True, color=PRIMARY_COLOR)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(p, after=2)
    r = p.add_run("14 questions pivots — État des lieux des outils de gestion")
    set_run_font(r, size=10, italic=True, color="555555")

    # ----- Identification compacte -----
    info = doc.add_table(rows=1, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    widths_info = [Cm(3.0), Cm(6.0), Cm(3.0), Cm(5.0)]
    info_cells = [
        ("Établissement :", "", "Date :", ""),
    ]
    for r_idx, row_data in enumerate(info_cells):
        row = info.rows[r_idx]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = widths_info[c_idx]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]
            set_paragraph_compact(cp)
            run = cp.add_run(val)
            if c_idx % 2 == 0:
                set_run_font(run, size=9, bold=True)
                set_cell_background(cell, LIGHT_FILL)
            else:
                set_run_font(run, size=9)

    info2 = doc.add_table(rows=1, cols=4)
    info2.alignment = WD_TABLE_ALIGNMENT.CENTER
    info2.autofit = False
    info2_cells = [
        ("Répondant :", "", "Fonction :", ""),
    ]
    for r_idx, row_data in enumerate(info2_cells):
        row = info2.rows[r_idx]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = widths_info[c_idx]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]
            set_paragraph_compact(cp)
            run = cp.add_run(val)
            if c_idx % 2 == 0:
                set_run_font(run, size=9, bold=True)
                set_cell_background(cell, LIGHT_FILL)
            else:
                set_run_font(run, size=9)

    # ----- Légende -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(p, before=4, after=4)
    r = p.add_run(
        "Pour chaque question, cochez la case qui décrit le mieux votre situation actuelle :   "
        "☐ Oui, totalement   ☐ Partiellement   ☐ Non"
    )
    set_run_font(r, size=9, italic=True, color="555555")

    # ----- Tableau des 14 questions -----
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(0.8), Cm(9.4), Cm(1.2), Cm(1.5), Cm(1.2), Cm(3.3)]
    headers = ["#", "Question pivot", "Oui", "Partiel.", "Non", "Observations"]

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_borders(cell)
        set_cell_background(cell, SECONDARY_COLOR)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = cell.paragraphs[0]
        set_paragraph_compact(cp)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(h)
        set_run_font(run, size=9, bold=True, color="FFFFFF")

    for i, (num, axe_titre, question) in enumerate(QUESTIONS):
        row = table.add_row()
        for c_idx in range(6):
            row.cells[c_idx].width = widths[c_idx]
            set_cell_borders(row.cells[c_idx])
            row.cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        alt = (i % 2 == 1)
        if alt:
            for c in row.cells:
                set_cell_background(c, LIGHT_FILL)

        # # axe
        cell = row.cells[0]
        cp = cell.paragraphs[0]
        set_paragraph_compact(cp)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(str(num))
        set_run_font(run, size=10, bold=True, color=PRIMARY_COLOR)

        # Question — titre de l'axe (gras) + question dessous
        cell = row.cells[1]
        cp = cell.paragraphs[0]
        set_paragraph_compact(cp)
        run = cp.add_run(axe_titre)
        set_run_font(run, size=8, bold=True, color=PRIMARY_COLOR)

        cp2 = cell.add_paragraph()
        set_paragraph_compact(cp2)
        run = cp2.add_run(question)
        set_run_font(run, size=9)

        # Cases à cocher
        for col_idx in (2, 3, 4):
            cell = row.cells[col_idx]
            cp = cell.paragraphs[0]
            set_paragraph_compact(cp)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run("☐")
            set_run_font(run, name="Segoe UI Symbol", size=14)

        # Observations
        cell = row.cells[5]
        cp = cell.paragraphs[0]
        set_paragraph_compact(cp)

    # ----- Synthèse intuitive -----
    p = doc.add_paragraph()
    set_paragraph_compact(p, before=8, after=2)
    r = p.add_run("Lecture rapide")
    set_run_font(r, size=10, bold=True, color=PRIMARY_COLOR)

    p = doc.add_paragraph()
    set_paragraph_compact(p, after=2)
    r = p.add_run("• Majorité de « Oui » : ")
    set_run_font(r, size=9, bold=True, color="00875A")
    r = p.add_run("système d'information mature, à optimiser à la marge.")
    set_run_font(r, size=9)

    p = doc.add_paragraph()
    set_paragraph_compact(p, after=2)
    r = p.add_run("• Majorité de « Partiellement » : ")
    set_run_font(r, size=9, bold=True, color="C77F00")
    r = p.add_run("outils en place mais hétérogènes — chantier d'intégration prioritaire.")
    set_run_font(r, size=9)

    p = doc.add_paragraph()
    set_paragraph_compact(p, after=2)
    r = p.add_run("• Majorité de « Non » : ")
    set_run_font(r, size=9, bold=True, color="B91C1C")
    r = p.add_run("rattrapage majeur nécessaire — risque opérationnel et concurrentiel élevé.")
    set_run_font(r, size=9)

    # ----- Signature -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_compact(p, before=10)
    r = p.add_run("Signature du répondant : _______________________________")
    set_run_font(r, size=9, italic=True, color="555555")

    doc.save(out_path)


def main():
    out = Path(__file__).resolve().parent / "audit_questionnaire_si_flash.docx"
    build_document(out)
    print(f"Fichier généré : {out}")


if __name__ == "__main__":
    main()
