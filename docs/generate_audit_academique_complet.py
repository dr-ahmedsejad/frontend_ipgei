"""
Génère le document d'audit académique & pédagogique complet (Word .docx)
en reprenant la structure du document "Commission de Transformation Digitale".

Ajouts intégrés dans la section C2 :
  - Compléments aux 3 sous-sections existantes (scolarité, pédagogie, évaluations)
  - 4 nouvelles sous-sections : vacations, réclamations, communication, sécurité

Sortie : docs/Audit_Academique_Pedagogique_complet.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Palette graphique (reprise du document original)
# ---------------------------------------------------------------------------

PRIMARY = "1F4E79"
LIGHT_BG = "F2F6FA"
ACCENT = "C9A227"
GREY_TEXT = "555555"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_run_font(run, name="Calibri", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def compact(p, before=0, after=0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0


def add_section_band(doc, label):
    """Bande de section principale (A., B., C1., etc.)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(17.5)
    set_cell_background(cell, PRIMARY)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    compact(p, before=2, after=2)
    run = p.add_run(f"  {label}")
    set_run_font(run, size=11, bold=True, color="FFFFFF")


def add_subsection_label(doc, label):
    """Sous-titre type B.1, C1.1, C2.1 — gras bleu avec léger trait."""
    p = doc.add_paragraph()
    compact(p, before=4, after=2)
    run = p.add_run(label)
    set_run_font(run, size=10, bold=True, color=PRIMARY)


def add_italic_note(doc, text):
    p = doc.add_paragraph()
    compact(p, before=2, after=2)
    run = p.add_run(text)
    set_run_font(run, size=9, italic=True, color=GREY_TEXT)


def add_checkbox_options(doc, options, single_column=True, cols=2):
    """Liste d'options avec case à cocher."""
    if single_column:
        for opt in options:
            t = doc.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = t.rows[0].cells[0]
            cell.width = Cm(17.5)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            compact(p)
            run = p.add_run("☐  ")
            set_run_font(run, name="Segoe UI Symbol", size=11)
            run = p.add_run(opt)
            set_run_font(run, size=10)
    else:
        rows = (len(options) + cols - 1) // cols
        t = doc.add_table(rows=rows, cols=cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        col_w = Cm(17.5 / cols)
        for r in range(rows):
            for c in range(cols):
                idx = r * cols + c
                cell = t.rows[r].cells[c]
                cell.width = col_w
                set_cell_borders(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                compact(p)
                if idx < len(options):
                    run = p.add_run("☐  ")
                    set_run_font(run, name="Segoe UI Symbol", size=11)
                    run = p.add_run(options[idx])
                    set_run_font(run, size=10)


def add_open_field(doc, label, lines=2):
    """Champ libre avec label en haut puis lignes de saisie."""
    p = doc.add_paragraph()
    compact(p, before=4, after=2)
    run = p.add_run(label)
    set_run_font(run, size=10, bold=True)
    for _ in range(lines):
        p = doc.add_paragraph()
        compact(p, before=2, after=2)
        run = p.add_run("_" * 110)
        set_run_font(run, size=9, color="BFBFBF")


def add_processus_table(doc, processus_list):
    """Tableau Processus / Aucune / Papier / Excel / Logiciel / Mixte / Priorité."""
    headers = ["Processus", "Aucune", "Papier", "Excel", "Logiciel", "Mixte", "Priorité"]
    widths = [Cm(7.5), Cm(1.4), Cm(1.4), Cm(1.4), Cm(1.4), Cm(1.4), Cm(3.0)]

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_borders(cell)
        set_cell_background(cell, PRIMARY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color="FFFFFF")

    for proc in processus_list:
        row = table.add_row()
        for c_idx in range(7):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Nom du processus
        cell = row.cells[0]
        p = cell.paragraphs[0]
        compact(p)
        run = p.add_run(proc)
        set_run_font(run, size=9, bold=True)

        # Cases situation (5 cases)
        for c_idx in range(1, 6):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            compact(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("☐")
            set_run_font(run, name="Segoe UI Symbol", size=12)

        # Priorité H / M / B
        cell = row.cells[6]
        p = cell.paragraphs[0]
        compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for prio in ("H", "M", "B"):
            run = p.add_run("☐")
            set_run_font(run, name="Segoe UI Symbol", size=11)
            run = p.add_run(prio + "  ")
            set_run_font(run, size=9, bold=True)


def add_subblock_title(doc, label):
    """Petit titre gras pour annoncer un bloc dans C2."""
    p = doc.add_paragraph()
    compact(p, before=6, after=2)
    run = p.add_run(label)
    set_run_font(run, size=10, bold=True, color=PRIMARY)


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

def build_document(out_path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

    # ===== En-tête =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    compact(p, after=0)
    run = p.add_run("Commission de Transformation Digitale  —  ")
    set_run_font(run, size=9, color=GREY_TEXT)
    run = p.add_run("Audit académique & pédagogique")
    set_run_font(run, size=9, bold=True, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact(p, before=2, after=0)
    run = p.add_run("COMMISSION DE TRANSFORMATION DIGITALE")
    set_run_font(run, size=11, bold=True, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact(p, after=4)
    run = p.add_run("Groupe d'Établissements d'Enseignement Supérieur — Mauritanie")
    set_run_font(run, size=10, italic=True, color=GREY_TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact(p, after=2)
    run = p.add_run("DOCUMENT D'AUDIT")
    set_run_font(run, size=18, bold=True, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact(p, after=2)
    run = p.add_run("Diagnostic & Recueil des Besoins — Volet Académique et Pédagogique")
    set_run_font(run, size=11, bold=True, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact(p, after=6)
    run = p.add_run("Périmètre : scolarité, enseignement, évaluations et examens")
    set_run_font(run, size=9, italic=True, color=GREY_TEXT)

    # Bandeau Établissement / Date
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    widths_h = [Cm(3.5), Cm(6.5), Cm(2.5), Cm(5.0)]
    labels_h = [("Établissement", True), ("", False), ("Date / Réf.", True), ("", False)]
    for c_idx, (val, is_label) in enumerate(labels_h):
        cell = t.rows[0].cells[c_idx]
        cell.width = widths_h[c_idx]
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if is_label:
            set_cell_background(cell, PRIMARY)
        p = cell.paragraphs[0]
        compact(p)
        run = p.add_run(val)
        if is_label:
            set_run_font(run, size=10, bold=True, color="FFFFFF")
        else:
            set_run_font(run, size=10)

    # ===== Section A — Identification =====
    add_section_band(doc, "A.  Identification de l'établissement")

    fields_A = [
        "Dénomination officielle",
        "Statut (public / privé / mixte) — Tutelle",
        "Nombre d'étudiants / d'enseignants",
        "Nombre de facultés ou départements — Sites",
        "Filières et niveaux proposés",
        "Langues d'enseignement",
    ]
    t = doc.add_table(rows=len(fields_A), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for r_idx, label in enumerate(fields_A):
        row = t.rows[r_idx]
        cell_label = row.cells[0]
        cell_label.width = Cm(6.0)
        set_cell_borders(cell_label)
        set_cell_background(cell_label, LIGHT_BG)
        cell_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell_label.paragraphs[0]
        compact(p)
        run = p.add_run(label)
        set_run_font(run, size=9, bold=True)

        cell_val = row.cells[1]
        cell_val.width = Cm(11.5)
        set_cell_borders(cell_val)
        compact(cell_val.paragraphs[0])

    # Interlocuteurs
    add_subsection_label(doc, "Interlocuteurs rencontrés (DG, DSI, chefs de service académiques)")

    headers_I = ["#", "Nom & Prénom", "Fonction", "Service", "Signature"]
    widths_I = [Cm(0.8), Cm(5.5), Cm(4.2), Cm(3.5), Cm(3.5)]
    t = doc.add_table(rows=5, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for c_idx, h in enumerate(headers_I):
        cell = t.rows[0].cells[c_idx]
        cell.width = widths_I[c_idx]
        set_cell_borders(cell)
        set_cell_background(cell, PRIMARY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color="FFFFFF")
    for r_idx in range(1, 5):
        for c_idx in range(5):
            cell = t.rows[r_idx].cells[c_idx]
            cell.width = widths_I[c_idx]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            compact(p)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(r_idx))
                set_run_font(run, size=9, bold=True)

    # ===== Section B — Diagnostic global =====
    add_section_band(doc, "B.  Diagnostic global de la digitalisation")
    add_italic_note(doc, "Section pivot : oriente la suite de l'audit (vers évaluation de l'existant, recueil des besoins, ou les deux).")

    add_subsection_label(doc, "B.1  Niveau global de digitalisation académique/pédagogique (une seule réponse) :")
    add_checkbox_options(doc, [
        "Aucune solution numérique — tout fonctionne sur papier",
        "Très partielle — quelques fichiers Excel/Word isolés",
        "Partielle — quelques logiciels en place, non intégrés",
        "Mixte — certains processus digitalisés, d'autres sur papier",
        "Avancée — la plupart des processus digitalisés, système fragmenté",
        "Intégrée — système d'information académique unifié",
    ])

    add_subsection_label(doc, "B.2  Plateformes pédagogiques en place :")
    add_checkbox_options(doc, [
        "LMS (Moodle, Google Classroom, autre)",
        "Outils de visioconférence (Zoom, Teams, Meet)",
        "ENT / Portail étudiant",
        "Application mobile étudiants",
        "Aucune plateforme pédagogique",
    ], single_column=False, cols=2)

    add_subsection_label(doc, "B.3  Inventaire des solutions numériques académiques en place")
    add_checkbox_options(doc, [
        "Aucune solution numérique en place — passer directement à la Section C2.",
    ])

    headers_inv = ["Nom de la solution", "Périmètre / Domaine couvert", "Éditeur / Fournisseur", "Année de mise en service"]
    widths_inv = [Cm(4.5), Cm(6.5), Cm(4.0), Cm(2.5)]
    t = doc.add_table(rows=5, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for c_idx, h in enumerate(headers_inv):
        cell = t.rows[0].cells[c_idx]
        cell.width = widths_inv[c_idx]
        set_cell_borders(cell)
        set_cell_background(cell, PRIMARY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color="FFFFFF")
    for r_idx in range(1, 5):
        for c_idx in range(4):
            cell = t.rows[r_idx].cells[c_idx]
            cell.width = widths_inv[c_idx]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            compact(cell.paragraphs[0])

    # ===== Section C1 — Évaluation des solutions existantes =====
    add_section_band(doc, "C1.  Évaluation des solutions existantes (si applicable)")
    add_italic_note(doc, "À renseigner uniquement si l'établissement dispose déjà de solutions. Objectif : conserver, faire évoluer ou remplacer.")

    headers_C1 = ["Critère d'évaluation", "1 Faible", "2", "3", "4", "5 Élevé", "N/A"]
    widths_C1 = [Cm(7.5), Cm(1.7), Cm(1.4), Cm(1.4), Cm(1.4), Cm(1.7), Cm(2.4)]
    criteres = [
        "Couverture fonctionnelle (besoin métier couvert)",
        "Performance, rapidité et fiabilité",
        "Ergonomie et facilité d'utilisation",
        "Sécurité des données et des accès",
        "Capacité d'évolution et interopérabilité",
        "Qualité du support et de la documentation",
    ]
    t = doc.add_table(rows=1 + len(criteres), cols=7)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for c_idx, h in enumerate(headers_C1):
        cell = t.rows[0].cells[c_idx]
        cell.width = widths_C1[c_idx]
        set_cell_borders(cell)
        set_cell_background(cell, PRIMARY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color="FFFFFF")
    for r_idx, crit in enumerate(criteres, start=1):
        row = t.rows[r_idx]
        for c_idx in range(7):
            cell = row.cells[c_idx]
            cell.width = widths_C1[c_idx]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell = row.cells[0]
        p = cell.paragraphs[0]
        compact(p)
        run = p.add_run(crit)
        set_run_font(run, size=9, bold=True)
        for c_idx in range(1, 7):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            compact(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("☐")
            set_run_font(run, name="Segoe UI Symbol", size=12)

    add_open_field(doc, "C1.1  Principales forces des solutions actuelles :")
    add_open_field(doc, "C1.2  Principales faiblesses ou dysfonctionnements :")

    add_subsection_label(doc, "C1.3  Position de l'établissement sur le devenir des solutions existantes :")
    add_checkbox_options(doc, [
        "À conserver en l'état — satisfaisantes",
        "À faire évoluer / compléter",
        "À remplacer progressivement",
        "À remplacer en urgence",
        "Arbitrage par la Commission",
    ])

    # ===== Section C2 — Recueil des besoins =====
    add_section_band(doc, "C2.  Recueil des besoins — Volet académique et pédagogique")
    add_italic_note(doc, "Pour chaque processus, indiquer la situation actuelle et la priorité de digitalisation (H = Haute, M = Moyenne, B = Basse).")

    # --- C2.A — Scolarité et gestion des étudiants (5 d'origine + 4 ajouts) ---
    add_subblock_title(doc, "Scolarité et gestion des étudiants")
    add_processus_table(doc, [
        "Inscriptions / Réinscriptions",
        "Dossiers étudiants",
        "Attestations & diplômes",
        "Emplois du temps & affectation des salles",
        "Suivi de l'assiduité / absences",
        # Ajouts SIGA
        "Pré-inscription / admission en ligne",
        "Portail étudiant (espace personnel : notes, EDT, documents)",
        "Suivi de la progression (passages, redoublements, exclusions)",
        "Justificatifs d'absence avec circuit de validation",
    ])

    # --- C2.B — Pédagogie et enseignement (5 d'origine + 3 ajouts) ---
    add_subblock_title(doc, "Pédagogie et enseignement")
    add_processus_table(doc, [
        "Programmes, syllabus et supports de cours",
        "Plateforme e-learning (LMS)",
        "Cours hybrides / visioconférence",
        "Travaux dirigés et travaux pratiques",
        "Suivi des stages et mémoires",
        # Ajouts SIGA
        "Pointage des séances effectivement tenues",
        "Suivi des charges horaires des enseignants",
        "Portail enseignant (espace personnel)",
    ])

    # --- C2.C — Évaluations et examens (5 d'origine + 3 ajouts) ---
    add_subblock_title(doc, "Évaluations et examens")
    add_processus_table(doc, [
        "Planification & convocations",
        "Saisie et calcul des notes",
        "Délibérations & jurys",
        "Publication des résultats",
        "Recours et archivage des copies",
        # Ajouts SIGA
        "Anonymat des copies à la saisie",
        "Sessions de rattrapage et compensation",
        "Pondérations CC / TP / Examen paramétrables",
    ])

    # --- C2.D — Vacations et paiements enseignants (nouveau bloc) ---
    add_subblock_title(doc, "Vacations et paiements enseignants")
    add_processus_table(doc, [
        "Calcul des vacations",
        "Génération des états de paiement",
        "Heures supplémentaires",
        "Export bancaire pour virements",
    ])

    # --- C2.E — Réclamations et vie étudiante (nouveau bloc) ---
    add_subblock_title(doc, "Réclamations et vie étudiante")
    add_processus_table(doc, [
        "Dépôt de réclamations en ligne (notes, absences)",
        "Périodes d'ouverture des réclamations",
        "Traitement et suivi du circuit",
    ])

    # --- C2.F — Communication et reporting (nouveau bloc) ---
    add_subblock_title(doc, "Communication et reporting")
    add_processus_table(doc, [
        "Notifications internes (intra-application)",
        "Notifications email / SMS sortantes",
        "Tableaux de bord direction (statistiques)",
        "Génération automatique d'attestations et relevés",
    ])

    # --- C2.G — Sécurité, traçabilité et multi-établissement (nouveau bloc) ---
    add_subblock_title(doc, "Sécurité, traçabilité et multi-établissement")
    add_processus_table(doc, [
        "Comptes individuels et rôles (direction, scolarité, enseignant, étudiant)",
        "Journal des modifications (qui a modifié quoi, quand)",
        "Gestion multi-établissement (isolation des données par site)",
    ])

    add_open_field(doc, "C2.1  3 besoins les plus critiques à couvrir en priorité (classer 1, 2, 3) :", lines=3)
    add_open_field(doc, "C2.2  Tâches manuelles à automatiser en priorité :", lines=3)

    # ===== Section D — Site web et présence en ligne =====
    add_section_band(doc, "D.  Site web et présence en ligne")
    add_italic_note(doc, "Vitrine numérique de l'établissement — communication externe, attractivité, relation candidats.")

    add_subsection_label(doc, "D.1  État du site web institutionnel (une seule réponse) :")
    add_checkbox_options(doc, [
        "Site web à jour — refonte de moins de 3 ans",
        "Site existant mais ancien — refonte de plus de 3 ans",
        "Site en cours de refonte",
        "Pas de site web institutionnel",
    ])

    add_subsection_label(doc, "D.2  Caractéristiques du site web (cocher tout ce qui s'applique) :")
    add_checkbox_options(doc, [
        "Responsive mobile",
        "Multilingue (FR / AR / EN)",
        "Catalogue des formations à jour",
        "Calendrier universitaire publié",
        "Annuaire enseignants / responsables consultable",
        "Formulaire de pré-candidature en ligne",
        "Lien vers portails (étudiant, enseignant)",
        "Mentions légales / politique RGPD affichées",
        "Statistiques de fréquentation suivies (Analytics)",
    ], single_column=False, cols=2)

    add_subsection_label(doc, "D.3  Présence sur les réseaux sociaux (cocher tout ce qui s'applique) :")
    add_checkbox_options(doc, [
        "Facebook",
        "LinkedIn",
        "Instagram",
        "YouTube",
        "TikTok",
        "X (Twitter)",
        "WhatsApp Business",
        "Aucune présence sur les réseaux sociaux",
    ], single_column=False, cols=3)

    add_subsection_label(doc, "D.4  Technologies utilisées (si applicable) :")

    # Mini-tableau technologies
    techs = [
        ("CMS / Plateforme utilisée", "(ex. WordPress, Drupal, Strapi, Next.js, développement sur mesure)"),
        ("Hébergement", "(prestataire / on-premise / cloud)"),
        ("Prestataire de maintenance / agence", "(nom, contact)"),
    ]
    t = doc.add_table(rows=len(techs), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for r_idx, (label, hint) in enumerate(techs):
        row = t.rows[r_idx]
        # Label
        cell_label = row.cells[0]
        cell_label.width = Cm(6.5)
        set_cell_borders(cell_label)
        set_cell_background(cell_label, LIGHT_BG)
        cell_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell_label.paragraphs[0]
        compact(p)
        run = p.add_run(label)
        set_run_font(run, size=9, bold=True)
        p2 = cell_label.add_paragraph()
        compact(p2)
        run = p2.add_run(hint)
        set_run_font(run, size=8, italic=True, color=GREY_TEXT)

        # Champ libre
        cell_val = row.cells[1]
        cell_val.width = Cm(11.0)
        set_cell_borders(cell_val)
        cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        compact(cell_val.paragraphs[0])

    # ===== Section E — Contraintes et environnement =====
    add_section_band(doc, "E.  Contraintes et environnement")

    add_subsection_label(doc, "E.1  Connexion Internet sur les sites :")
    add_checkbox_options(doc, [
        "Très bonne / bonne",
        "Moyenne — coupures fréquentes",
        "Faible / pas de connexion sur certains sites",
    ])

    add_subsection_label(doc, "E.2  Compétences IT internes :")
    add_checkbox_options(doc, [
        "DSI structurée",
        "Quelques techniciens",
        "Une seule personne en charge",
        "Pas de compétences internes",
    ], single_column=False, cols=2)

    add_subsection_label(doc, "E.3  Préférences linguistiques pour les interfaces :")
    add_checkbox_options(doc, [
        "Français uniquement",
        "Bilingue français / arabe",
        "Trilingue français / arabe / anglais",
    ], single_column=False, cols=3)

    add_open_field(doc, "E.4  Budget prévisionnel ou sources de financement identifiées :")
    add_open_field(doc, "E.5  Délai souhaité pour les premiers résultats / déploiement complet :")

    # ===== Section F — Synthèse =====
    add_section_band(doc, "F.  Synthèse — Réservé à la Commission")
    add_italic_note(doc, "À renseigner par la Commission à l'issue de l'audit. Base de la recommandation au Groupe.")

    add_subsection_label(doc, "F.1  Niveau de maturité digitale constaté (volet académique/pédagogique) :")
    add_checkbox_options(doc, [
        "Niveau 1 — Inexistant (tout manuel)",
        "Niveau 2 — Émergent (outils bureautiques isolés)",
        "Niveau 3 — En construction (logiciels partiels)",
        "Niveau 4 — Structuré (système partiellement intégré)",
        "Niveau 5 — Mature (SI académique unifié)",
    ])

    add_subsection_label(doc, "F.2  Orientation recommandée :")
    add_checkbox_options(doc, [
        "Nouvelle solution complète à concevoir",
        "Évolution / complétion des solutions existantes",
        "Refonte partielle (remplacement de modules)",
        "Mutualisation avec d'autres établissements du Groupe",
        "Mise à niveau préalable de l'infrastructure",
        "Étude complémentaire nécessaire",
    ])

    add_open_field(doc, "F.3  Axes prioritaires retenus et recommandations :", lines=4)

    # ===== Section G — Validation et signatures =====
    add_section_band(doc, "G.  Validation et signatures")
    add_italic_note(doc, "Les soussignés certifient que les informations renseignées dans le présent document reflètent fidèlement la situation de l'établissement à la date d'audit.")

    headers_F = ["POUR L'ÉTABLISSEMENT", "POUR LA COMMISSION"]
    t = doc.add_table(rows=4, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for c_idx, h in enumerate(headers_F):
        cell = t.rows[0].cells[c_idx]
        cell.width = Cm(8.7)
        set_cell_borders(cell)
        set_cell_background(cell, PRIMARY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        compact(p)
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color="FFFFFF")

    rows_F_left = ["Nom, Fonction :", "Date :", "Signature & cachet :"]
    rows_F_right = ["Président / Rapporteur :", "Date :", "Signatures :"]
    for r_idx, (l, r) in enumerate(zip(rows_F_left, rows_F_right), start=1):
        for c_idx, val in enumerate([l, r]):
            cell = t.rows[r_idx].cells[c_idx]
            cell.width = Cm(8.7)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            compact(p, before=4, after=12)
            run = p.add_run(val)
            set_run_font(run, size=9, bold=True)

    doc.save(out_path)


def main():
    out = Path(__file__).resolve().parent / "Audit_Academique_Pedagogique_complet.docx"
    try:
        build_document(out)
        print(f"Fichier généré : {out}")
    except PermissionError:
        alt = Path(__file__).resolve().parent / "Audit_Academique_Pedagogique_complet_v2.docx"
        build_document(alt)
        print(
            f"Fichier verrouillé (sans doute ouvert dans Word). "
            f"Sauvegardé sous : {alt}"
        )


if __name__ == "__main__":
    main()
