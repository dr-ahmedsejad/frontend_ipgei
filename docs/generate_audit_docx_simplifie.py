"""
Génère un questionnaire d'audit SI au format Word (.docx) — VERSION SIMPLIFIÉE.

Cible : directeur d'établissement, directeur des études, chef de département,
responsable scolarité (profils NON-techniques).

Caractéristiques :
- Langage accessible (pas de jargon IT)
- Cases à cocher : Oui / Partiellement / Non
- Maximum 5 pages
- ~60 questions sur 14 axes

Sortie : docs/audit_questionnaire_si_simplifie.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Contenu — formulation simplifiée, langage non-IT
# ---------------------------------------------------------------------------

AXES = [
    {
        "num": 1,
        "titre": "Pilotage informatique de l'établissement",
        "questions": [
            "Une personne est désignée pour piloter l'informatique de l'établissement (DSI, responsable IT ou prestataire) ?",
            "Existe-t-il un plan ou une feuille de route pour l'informatique de l'établissement ?",
            "Un budget annuel est-il prévu pour les outils numériques ?",
        ],
    },
    {
        "num": 2,
        "titre": "Sécurité des données",
        "questions": [
            "Les données sont-elles sauvegardées régulièrement (quotidien / hebdomadaire) ?",
            "Une sauvegarde est-elle stockée à l'extérieur de l'établissement (cloud, autre site) ?",
            "Chaque utilisateur a-t-il son propre compte (pas de compte partagé) ?",
            "Les droits d'accès sont-ils définis selon les fonctions (direction, scolarité, enseignant, étudiant) ?",
            "Un mécanisme existe-t-il pour bloquer / débloquer un compte ?",
            "L'établissement est-il en conformité avec la réglementation sur les données personnelles ?",
        ],
    },
    {
        "num": 3,
        "titre": "Traçabilité des actions",
        "questions": [
            "Les modifications importantes (notes, inscriptions…) sont-elles enregistrées avec l'auteur et la date ?",
            "Peut-on consulter l'historique d'un dossier étudiant ou d'une note ?",
            "Existe-t-il une protection contre la suppression ou la modification frauduleuse de l'historique ?",
        ],
    },
    {
        "num": 4,
        "titre": "Référentiels académiques",
        "questions": [
            "Le catalogue des filières, départements et modules est-il géré dans un outil informatique ?",
            "Le programme est-il rattaché à une année universitaire (versioning par année) ?",
            "Les semestres, semaines et créneaux horaires sont-ils paramétrables ?",
            "Le référentiel des salles existe-t-il (capacité, équipement) ?",
            "Le calendrier prend-il en compte les périodes spécifiques (Ramadan, jours fériés locaux) ?",
        ],
    },
    {
        "num": 5,
        "titre": "Pré-inscriptions et admissions",
        "questions": [
            "Les candidats peuvent-ils pré-déposer leur dossier en ligne (formulaire web) ?",
            "Chaque candidat reçoit-il un lien personnel sécurisé pour suivre sa candidature ?",
            "Le passage de la candidature à l'inscription définitive est-il géré dans le même outil ?",
            "L'import en masse d'étudiants (depuis un fichier Excel) est-il possible ?",
        ],
    },
    {
        "num": 6,
        "titre": "Vie étudiante et scolarité",
        "questions": [
            "Le dossier de chaque étudiant est-il centralisé (état civil, parcours, documents) ?",
            "Peut-on rechercher un étudiant selon plusieurs critères ?",
            "La progression de l'étudiant est-elle suivie d'une année / d'un semestre à l'autre ?",
            "Les étudiants ont-ils un espace personnel (notes, EDT, absences, documents) ?",
        ],
    },
    {
        "num": 7,
        "titre": "Emplois du temps",
        "questions": [
            "Les emplois du temps sont-ils conçus dans un outil dédié (pas Excel) ?",
            "L'outil détecte-t-il automatiquement les conflits (enseignant, salle, groupe occupés en même temps) ?",
            "Les séances de rattrapage et modifications de dernière minute sont-elles gérées ?",
            "Les emplois du temps sont-ils archivés par semaine / année ?",
            "Les enseignants et étudiants peuvent-ils consulter leur EDT en ligne ?",
        ],
    },
    {
        "num": 8,
        "titre": "Enseignants et paiement des vacations",
        "questions": [
            "L'annuaire des enseignants distingue-t-il permanents / vacataires / contractuels ?",
            "Les charges horaires de chaque enseignant sont-elles suivies dans le système ?",
            "Les états de paiement des vacations sont-ils générés automatiquement ?",
            "Les heures supplémentaires sont-elles calculées automatiquement ?",
            "Un export bancaire pour les virements existe-t-il ?",
        ],
    },
    {
        "num": 9,
        "titre": "Présences et absences",
        "questions": [
            "Le pointage de présence des enseignants est-il fait (séance effectivement tenue) ?",
            "Les absences des étudiants sont-elles saisies par séance / par salle ?",
            "Les justificatifs d'absence (médicaux, administratifs) sont-ils gérés avec un circuit de validation ?",
            "Les enseignants peuvent-ils saisir les absences depuis leur espace personnel ?",
            "Des alertes automatiques sont-elles envoyées au-delà d'un seuil d'absences ?",
        ],
    },
    {
        "num": 10,
        "titre": "Notes, évaluations et délibérations",
        "questions": [
            "Les enseignants saisissent-ils les notes depuis leur espace personnel ?",
            "La saisie anonyme des copies (anonymat préservé) est-elle possible ?",
            "Les sessions normale et de rattrapage sont-elles distinguées ?",
            "Les pondérations CC / TP / Examen sont-elles paramétrables par module ?",
            "Les moyennes (module, semestre, année) sont-elles calculées automatiquement ?",
            "Les procès-verbaux de délibération sont-ils générés automatiquement ?",
            "Les décisions du jury (rachat, passage, redoublement) sont-elles tracées ?",
        ],
    },
    {
        "num": 11,
        "titre": "Stages et insertion professionnelle",
        "questions": [
            "Les conventions de stage sont-elles gérées dans un outil informatique ?",
            "Le circuit de validation d'une convention est-il suivi (brouillon → soumise → terminée) ?",
            "La distinction stage / Projet de Fin d'Études (PFE) existe-t-elle ?",
            "Les tuteurs (académique + entreprise) sont-ils enregistrés ?",
            "L'évaluation finale prend-elle en compte plusieurs notes (entreprise, rapport, soutenance) ?",
            "Le jury de soutenance est-il composé de plusieurs membres tracés dans l'outil ?",
        ],
    },
    {
        "num": 12,
        "titre": "Réclamations étudiantes",
        "questions": [
            "Les étudiants peuvent-ils déposer une réclamation (absence, note, autre) depuis leur espace ?",
            "Les périodes d'ouverture des réclamations sont-elles paramétrables ?",
            "La réclamation est-elle reliée à l'objet concerné (séance, note, session) ?",
            "Le traitement suit-il un circuit (soumise → en cours → acceptée / rejetée) ?",
            "Les enseignants concernés voient-ils les réclamations qui les concernent ?",
        ],
    },
    {
        "num": 13,
        "titre": "Documents, communication et tableaux de bord",
        "questions": [
            "Les attestations, relevés de notes et autres documents sont-ils générés automatiquement ?",
            "Les étudiants peuvent-ils télécharger leurs documents depuis leur espace personnel ?",
            "Des notifications internes sont-elles envoyées aux utilisateurs ?",
            "Les notifications par e-mail ou SMS sont-elles possibles ?",
            "La direction dispose-t-elle de tableaux de bord (statistiques globales, par filière, par enseignant) ?",
            "Les exports pour la tutelle / ministère sont-ils possibles ?",
            "La signature électronique des documents officiels est-elle utilisée ?",
        ],
    },
    {
        "num": 14,
        "titre": "Site web et présence numérique",
        "questions": [
            "Le site web institutionnel est-il à jour (refonte récente, contenu actualisé) ?",
            "Le site est-il consultable correctement sur smartphone ?",
            "Le site est-il disponible en plusieurs langues (FR / AR / EN selon contexte) ?",
            "Le catalogue des formations en ligne est-il à jour ?",
            "Un formulaire de pré-candidature est-il accessible directement depuis le site ?",
            "L'établissement est-il présent sur les réseaux sociaux (Facebook, LinkedIn) ?",
            "Les statistiques de fréquentation du site sont-elles suivies ?",
            "Les mentions légales et politique de confidentialité (RGPD) sont-elles affichées ?",
        ],
    },
]


# ---------------------------------------------------------------------------
# Helpers de style
# ---------------------------------------------------------------------------

PRIMARY_COLOR = "1F4E79"
SECONDARY_COLOR = "2E75B6"
LIGHT_FILL = "F2F6FA"


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_run_font(run, name="Calibri", size=9, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_compact(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def add_axe_header(doc, num, titre):
    p = doc.add_paragraph()
    set_paragraph_compact(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Axe {num} — {titre}")
    set_run_font(run, size=10, bold=True, color="FFFFFF")
    # Fond coloré
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), PRIMARY_COLOR)
    p_pr.append(shd)


# ---------------------------------------------------------------------------
# Construction du document
# ---------------------------------------------------------------------------

def build_document(out_path: Path):
    doc = Document()

    # Marges réduites pour gagner de l'espace
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ----- Titre principal -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(title)
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Questionnaire d'audit du Système d'Information")
    set_run_font(r, size=14, bold=True, color=PRIMARY_COLOR)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(subtitle)
    subtitle.paragraph_format.space_after = Pt(4)
    r = subtitle.add_run("État des lieux des outils de gestion — Établissements supérieurs")
    set_run_font(r, size=10, color="555555")
    r.italic = True

    # ----- Encadré identification (compact) -----
    info_table = doc.add_table(rows=2, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False

    widths_info = [Cm(3.2), Cm(5.5), Cm(3.2), Cm(5.5)]
    info_fields = [
        ("Établissement :", "", "Date de l'audit :", ""),
        ("Répondant :", "", "Fonction :", ""),
    ]
    for r_idx, row_data in enumerate(info_fields):
        row = info_table.rows[r_idx]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = widths_info[c_idx]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph_compact(p)
            run = p.add_run(val)
            if c_idx % 2 == 0:
                set_run_font(run, size=9, bold=True)
                set_cell_background(cell, LIGHT_FILL)
            else:
                set_run_font(run, size=9)

    # ----- Légende -----
    legend = doc.add_paragraph()
    set_paragraph_compact(legend)
    legend.paragraph_format.space_before = Pt(4)
    legend.paragraph_format.space_after = Pt(2)
    legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = legend.add_run(
        "Cochez pour chaque question :  ☐ Oui   ☐ Partiellement   ☐ Non   |   Observations en colonne libre"
    )
    set_run_font(r, size=9, color="555555")
    r.italic = True

    # ----- Tableau questionnaire global -----
    # Une seule grande table : 5 colonnes — # / Question / Oui / Partiel / Non / Observations
    # Pour rester compact : on regroupe Oui/Partiel/Non en 3 colonnes étroites
    nb_rows = 1  # header (axe header en row span)
    for axe in AXES:
        nb_rows += 1  # ligne de titre d'axe
        nb_rows += len(axe["questions"])

    # On crée l'en-tête comme première ligne fixe
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Largeurs (total ~17.5 cm pour rester dans les marges)
    widths = [Cm(0.9), Cm(9.5), Cm(1.0), Cm(1.4), Cm(1.0), Cm(3.7)]
    headers = ["#", "Question", "Oui", "Partiel.", "Non", "Observations"]

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_borders(cell)
        set_cell_background(cell, SECONDARY_COLOR)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph_compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color="FFFFFF")

    # Lignes : pour chaque axe, ajouter une ligne fusionnée puis les questions
    q_counter_per_axe = {}
    for axe in AXES:
        # Ligne axe (fusion 6 cols)
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            set_cell_borders(cell)
        merged = row.cells[0].merge(row.cells[5])
        set_cell_background(merged, PRIMARY_COLOR)
        merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = merged.paragraphs[0]
        set_paragraph_compact(p)
        run = p.add_run(f"  Axe {axe['num']} — {axe['titre']}")
        set_run_font(run, size=10, bold=True, color="FFFFFF")

        # Lignes des questions
        for q_idx, question in enumerate(axe["questions"], start=1):
            num = f"{axe['num']}.{q_idx}"
            q_counter_per_axe[axe['num']] = q_idx
            row = table.add_row()
            row.height = Cm(0.45)

            # # numéro
            cell = row.cells[0]
            cell.width = widths[0]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph_compact(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(num)
            set_run_font(run, size=8, bold=True)

            # Question
            cell = row.cells[1]
            cell.width = widths[1]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph_compact(p)
            run = p.add_run(question)
            set_run_font(run, size=9)

            # Oui / Partiel / Non — case à cocher centrée
            for col_idx in (2, 3, 4):
                cell = row.cells[col_idx]
                cell.width = widths[col_idx]
                set_cell_borders(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                set_paragraph_compact(p)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("☐")
                set_run_font(run, name="Segoe UI Symbol", size=12)

            # Observations
            cell = row.cells[5]
            cell.width = widths[5]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph_compact(p)

            # Alternance de fond (rangées paires plus claires)
            if q_idx % 2 == 0:
                for c in row.cells:
                    set_cell_background(c, LIGHT_FILL)

    # ----- Pied de page : signature -----
    sig_para = doc.add_paragraph()
    set_paragraph_compact(sig_para)
    sig_para.paragraph_format.space_before = Pt(8)
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig_para.add_run("Signature du répondant : _______________________________")
    set_run_font(r, size=9, color="555555")
    r.italic = True

    doc.save(out_path)


def main():
    out = Path(__file__).resolve().parent / "audit_questionnaire_si_simplifie.docx"
    build_document(out)
    print(f"Fichier généré : {out}")


if __name__ == "__main__":
    main()
