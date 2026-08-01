"""
Génère un questionnaire d'audit SI au format Word (.docx) — VERSION 30 QUESTIONS.

30 questions directes (sans scénario), couvrant les 14 axes du questionnaire PDF.
Format compact 2-3 pages avec cases à cocher Oui / Partiellement / Non.

Sortie : docs/audit_questionnaire_si_30q.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# 30 questions directes, regroupées par axe (14 axes)
# ---------------------------------------------------------------------------

AXES = [
    (
        1, "Pilotage informatique",
        [
            "Une personne ou une équipe est-elle désignée pour piloter le système d'information ?",
            "Un budget annuel est-il prévu pour les outils informatiques de gestion ?",
        ],
    ),
    (
        2, "Sécurité des données",
        [
            "Les données sont-elles sauvegardées quotidiennement avec une copie externalisée ?",
            "Chaque utilisateur dispose-t-il d'un compte personnel avec mot de passe individuel ?",
            "Les droits d'accès sont-ils définis par profil (direction, scolarité, enseignant, étudiant) ?",
        ],
    ),
    (
        3, "Traçabilité des actions",
        [
            "L'historique des modifications (notes, inscriptions, dossiers) est-il enregistré avec auteur et date ?",
            "Les logs / journaux d'activité sont-ils consultables et exportables par les responsables ?",
        ],
    ),
    (
        4, "Référentiels académiques",
        [
            "Les filières, modules, salles, semestres et calendriers sont-ils gérés dans un seul système ?",
            "Le programme pédagogique est-il rattaché à une année universitaire (versioning) ?",
        ],
    ),
    (
        5, "Pré-inscriptions et admissions",
        [
            "Les candidats peuvent-ils déposer leur dossier de pré-inscription en ligne via un formulaire web ?",
            "L'import en masse d'étudiants (fichier Excel / CSV) est-il possible ?",
        ],
    ),
    (
        6, "Vie étudiante et scolarité",
        [
            "Le dossier de chaque étudiant est-il centralisé dans le système (état civil, parcours, documents) ?",
            "Les étudiants ont-ils un espace personnel en ligne (notes, EDT, absences, documents) ?",
        ],
    ),
    (
        7, "Emplois du temps",
        [
            "Les emplois du temps sont-ils générés dans un outil dédié (et non sur Excel) ?",
            "L'outil détecte-t-il automatiquement les conflits (enseignant, salle, groupe) ?",
        ],
    ),
    (
        8, "Enseignants et vacations",
        [
            "L'annuaire des enseignants distingue-t-il permanents, vacataires et contractuels ?",
            "Les états de paiement des vacations et heures supplémentaires sont-ils générés automatiquement ?",
        ],
    ),
    (
        9, "Présences et absences",
        [
            "Le pointage des séances et la saisie des absences sont-ils effectués dans le système ?",
            "Les justificatifs d'absence sont-ils gérés avec un circuit de validation ?",
        ],
    ),
    (
        10, "Évaluations et délibérations",
        [
            "Les enseignants saisissent-ils les notes depuis leur espace personnel en ligne ?",
            "La saisie anonyme des copies (anonymat préservé) est-elle possible ?",
            "Les procès-verbaux de délibération sont-ils générés automatiquement (moyennes, décisions, rachats) ?",
        ],
    ),
    (
        11, "Stages et insertion professionnelle",
        [
            "Les conventions de stage sont-elles gérées dans le système avec workflow complet (brouillon → terminée) ?",
            "L'évaluation finale combine-t-elle plusieurs notes (entreprise, rapport, soutenance) ?",
        ],
    ),
    (
        12, "Réclamations étudiantes",
        [
            "Les étudiants peuvent-ils déposer une réclamation (note, absence, autre) en ligne ?",
            "Les périodes d'ouverture sont-elles paramétrables avec workflow de traitement tracé ?",
        ],
    ),
    (
        13, "Documents et reporting",
        [
            "Les attestations, relevés de notes et autres documents officiels sont-ils générés automatiquement ?",
            "La direction dispose-t-elle de tableaux de bord (statistiques par filière, par enseignant, globales) ?",
        ],
    ),
    (
        14, "Site web et présence numérique",
        [
            "Le site web institutionnel est-il à jour, responsive (mobile) et multilingue ?",
            "Le formulaire de pré-candidature et le catalogue des formations sont-ils accessibles en ligne ?",
        ],
    ),
]


# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

PRIMARY_COLOR = "1F4E79"
SECONDARY_COLOR = "2E75B6"
LIGHT_FILL = "F2F6FA"


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def set_run_font(run, name="Calibri", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_compact(p, before=0, after=0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

def build_document(out_path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

    # ----- Titre -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(p, after=2)
    r = p.add_run("Questionnaire d'audit du Système d'Information")
    set_run_font(r, size=14, bold=True, color=PRIMARY_COLOR)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(p, after=3)
    r = p.add_run("30 questions essentielles — État des lieux des outils de gestion")
    set_run_font(r, size=10, italic=True, color="555555")

    # ----- Identification compacte (2 lignes) -----
    for fields in [
        ("Établissement :", "", "Date :", ""),
        ("Répondant :", "", "Fonction :", ""),
    ]:
        info = doc.add_table(rows=1, cols=4)
        info.alignment = WD_TABLE_ALIGNMENT.CENTER
        info.autofit = False
        widths_info = [Cm(3.0), Cm(6.5), Cm(2.5), Cm(5.0)]
        row = info.rows[0]
        for c_idx, val in enumerate(fields):
            cell = row.cells[c_idx]
            cell.width = widths_info[c_idx]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cp = cell.paragraphs[0]
            set_paragraph_compact(cp)
            run = cp.add_run(val)
            if c_idx % 2 == 0:
                set_run_font(run, size=9, bold=True)
                set_cell_background(cell, LIGHT_FILL)
            else:
                set_run_font(run, size=9)

    # ----- Légende -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_compact(p, before=4, after=4)
    r = p.add_run(
        "Cochez la case qui décrit votre situation :   "
        "☐ Oui   ☐ Partiellement   ☐ Non    |    Précisions en colonne « Observations »"
    )
    set_run_font(r, size=9, italic=True, color="555555")

    # ----- Tableau des questions -----
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(0.9), Cm(9.4), Cm(1.0), Cm(1.5), Cm(1.0), Cm(3.5)]
    headers = ["#", "Question", "Oui", "Partiel.", "Non", "Observations"]

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_borders(cell)
        set_cell_background(cell, SECONDARY_COLOR)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = cell.paragraphs[0]
        set_paragraph_compact(cp)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(h)
        set_run_font(run, size=9, bold=True, color="FFFFFF")

    for axe_num, axe_titre, questions in AXES:
        # Ligne axe (fusion 6 cols)
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            set_cell_borders(cell)
        merged = row.cells[0].merge(row.cells[5])
        set_cell_background(merged, PRIMARY_COLOR)
        merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = merged.paragraphs[0]
        set_paragraph_compact(cp)
        run = cp.add_run(f"  Axe {axe_num} — {axe_titre}")
        set_run_font(run, size=10, bold=True, color="FFFFFF")

        # Lignes des questions
        for q_idx, question in enumerate(questions, start=1):
            num = f"{axe_num}.{q_idx}"
            row = table.add_row()
            row.height = Cm(0.55)

            for c_idx in range(6):
                row.cells[c_idx].width = widths[c_idx]
                set_cell_borders(row.cells[c_idx])
                row.cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            alt = (q_idx % 2 == 0)
            if alt:
                for c in row.cells:
                    set_cell_background(c, LIGHT_FILL)

            # # numéro
            cell = row.cells[0]
            cp = cell.paragraphs[0]
            set_paragraph_compact(cp)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run(num)
            set_run_font(run, size=9, bold=True, color=PRIMARY_COLOR)

            # Question
            cell = row.cells[1]
            cp = cell.paragraphs[0]
            set_paragraph_compact(cp)
            run = cp.add_run(question)
            set_run_font(run, size=10)

            # Cases à cocher
            for col_idx in (2, 3, 4):
                cell = row.cells[col_idx]
                cp = cell.paragraphs[0]
                set_paragraph_compact(cp)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cp.add_run("☐")
                set_run_font(run, name="Segoe UI Symbol", size=14)

            # Observations (vide)
            cell = row.cells[5]
            cp = cell.paragraphs[0]
            set_paragraph_compact(cp)

    # ----- Lecture rapide -----
    p = doc.add_paragraph()
    set_paragraph_compact(p, before=8, after=2)
    r = p.add_run("Lecture rapide du résultat")
    set_run_font(r, size=10, bold=True, color=PRIMARY_COLOR)

    p = doc.add_paragraph()
    set_paragraph_compact(p, after=2)
    r = p.add_run("• Majorité de « Oui » : ")
    set_run_font(r, size=9, bold=True, color="00875A")
    r = p.add_run("système d'information mature, à optimiser à la marge.")
    set_run_font(r, size=9)

    p = doc.add_paragraph()
    set_paragraph_compact(p, after=2)
    r = p.add_run("• Majorité de « Partiellement » : ")
    set_run_font(r, size=9, bold=True, color="C77F00")
    r = p.add_run("outils en place mais hétérogènes — chantier d'intégration prioritaire.")
    set_run_font(r, size=9)

    p = doc.add_paragraph()
    set_paragraph_compact(p, after=2)
    r = p.add_run("• Majorité de « Non » : ")
    set_run_font(r, size=9, bold=True, color="B91C1C")
    r = p.add_run("rattrapage majeur nécessaire — risque opérationnel et concurrentiel élevé.")
    set_run_font(r, size=9)

    # ----- Signature -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_compact(p, before=8)
    r = p.add_run("Signature du répondant : _______________________________")
    set_run_font(r, size=9, italic=True, color="555555")

    doc.save(out_path)


def main():
    out = Path(__file__).resolve().parent / "audit_questionnaire_si_30q.docx"
    build_document(out)
    print(f"Fichier généré : {out}")


if __name__ == "__main__":
    main()
